from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter
import asyncio
import os
import logging
from pathlib import Path
from typing import Dict
from datetime import datetime

logger = logging.getLogger(__name__)


def safe_add(a, b):
    if a is None or b is None:
        return None
    return a + b

class DocumentGenerator:
    def __init__(self):
        self.company_name = None
        self.logo_path = None
        self.company_details: Dict[str, str] = {}

    def set_branding(self, company_name: str = None, logo_path: str = None):
        if company_name:
            self.company_name = company_name
        if logo_path:
            self.logo_path = logo_path
    
    async def generate_word(self, tender_data: Dict, pricing: Dict, output_path: str):
        loop = asyncio.get_event_loop()
        await loop.run_in_executor(None, self._create_professional_doc, tender_data, pricing, output_path)
    
    def _create_professional_doc(self, tender_data: Dict, pricing: Dict, output_path: str):
        doc = Document()
        
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # COVER PAGE
        self._add_cover_page(doc, tender_data, pricing)
        doc.add_page_break()
        
        # TABLE OF CONTENTS
        self._add_toc(doc)
        doc.add_page_break()
        
        # COMPANY PROFILE
        self._add_section_heading(doc, "1. COMPANY PROFILE")
        self._add_company_profile(doc)
        
        # UNDERSTANDING
        self._add_section_heading(doc, "2. UNDERSTANDING OF REQUIREMENTS")
        self._add_requirements_understanding(doc, tender_data)
        
        # TECHNICAL PROPOSAL
        self._add_section_heading(doc, "3. TECHNICAL PROPOSAL")
        self._add_technical_proposal(doc, tender_data, pricing)
        
        # PRICING
        self._add_section_heading(doc, "4. PRICING SCHEDULE")
        self._add_pricing_schedule(doc, pricing)
        
        # COMPLIANCE
        self._add_section_heading(doc, "5. COMPLIANCE CHECKLIST")
        self._add_compliance_checklist(doc)
        
        # DECLARATION
        self._add_section_heading(doc, "6. DECLARATION")
        self._add_declaration(doc)
        
        doc.save(output_path)
    
    def _add_cover_page(self, doc, tender_data, pricing):
        if self.logo_path and Path(self.logo_path).exists():
            try:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(self.logo_path, width=Inches(2.5))
            except Exception as e:
                logger.exception("Logo insertion error: %s", e)
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run("[COMPANY LOGO]")
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(128, 128, 128)
        else:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run("[COMPANY LOGO]")
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("TENDER SUBMISSION")
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        
        doc.add_paragraph()
        
        # Tender details
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        details = [
            ("Tender Reference:", tender_data.get('reference')),
            ("Tender Type:", (tender_data.get('sector') or tender_data.get('tender_type') or None).title() if (tender_data.get('sector') or tender_data.get('tender_type')) else None),
            ("Submission Date:", datetime.now().strftime('%d %B %Y')),
            ("Contract Duration:", (tender_data.get('duration') or {}).get('display') or (f"{tender_data.get('duration_months')} months" if tender_data.get('duration_months') is not None else None))
        ]
        
        for i, (label, value) in enumerate(details):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value) if value is not None else ''
            row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # CRITICAL FIX: pricing is now properly passed and used
        total_value = pricing.get('total_contract_value', 0) if isinstance(pricing, dict) else 0
        
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"Total Tender Value: R {total_value:,.2f}")
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 51)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run("Submitted by:").font.size = Pt(12)

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(self.company_name or '')
        run.font.size = Pt(14)
        run.font.bold = True
    
    def _add_toc(self, doc):
        para = doc.add_paragraph()
        run = para.add_run("TABLE OF CONTENTS")
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        doc.add_paragraph()
        
        toc_items = [
            "1. Company Profile",
            "2. Understanding of Requirements", 
            "3. Technical Proposal",
            "4. Pricing Schedule",
            "5. Compliance Checklist",
            "6. Declaration"
        ]
        
        for item in toc_items:
            doc.add_paragraph(item, style='List Number')
    
    def _add_section_heading(self, doc, text):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        para.space_after = Pt(12)
        doc.add_paragraph()
    
    def _add_company_profile(self, doc):
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        
        company_info = [
            ("Company Registration:", self.company_details.get('reg', '')),
            ("VAT Number:", self.company_details.get('vat', '')),
            ("CIDB Grade:", self.company_details.get('cidb', '')),
            ("B-BBEE Level:", self.company_details.get('bbbee', ''))
        ]
        
        for i, (label, value) in enumerate(company_info):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            row.cells[0].paragraphs[0].runs[0].font.bold = True
    
    def _add_requirements_understanding(self, doc, tender_data):
        req = tender_data.get('requirements', {})
        workforce = tender_data.get('workforce', {}) if isinstance(tender_data.get('workforce', {}), dict) else {}
        duration_months = tender_data.get('duration_months') or (tender_data.get('duration') or {}).get('months')
        sector = tender_data.get('sector') or tender_data.get('tender_type')
        workers = workforce.get('total_workers')
        supervisors = workforce.get('supervisors')
        if workers is None:
            workers = req.get('workers')
        if supervisors is None:
            supervisors = req.get('supervisors')

        required_fields = ["workers", "supervisors"]
        missing = [
            field for field in required_fields
            if tender_data.get('requirements', {}).get(field) is None
            and workforce.get('total_workers' if field == 'workers' else field) is None
        ]
        if missing:
            logger.warning("Missing optional fields: %s", missing)

        total_staff = safe_add(workers, supervisors)
        staff_complement = (
            f"Staff Complement: {total_staff} personnel"
            if total_staff is not None
            else "Staff Complement: Not specified in tender"
        )

        # Use executive summary if available
        exec_summary = tender_data.get('executive_summary')
        if exec_summary:
            doc.add_paragraph(exec_summary)
            doc.add_paragraph()
        
        # Use scope description if available
        scope_desc = tender_data.get('scope_description')
        if scope_desc:
            doc.add_paragraph("Scope of Work:", style='Heading 2')
            doc.add_paragraph(scope_desc)
            doc.add_paragraph()

        content = f"""We understand that this tender requires the provision of {sector or ''} services for a period of {duration_months if duration_months is not None else 'Not specified'} months at {tender_data.get('location') or 'the specified location'}."""
        
        doc.add_paragraph(content)
        doc.add_paragraph("\nKey Requirements Identified:")
        
        hours_value = req.get('hours_per_day')
        shifts_value = req.get('shifts_per_day', req.get('shifts'))
        area_value = tender_data.get('scope', {}).get('area_sqm') or req.get('area_sqm')

        requirements_list = [
            staff_complement,
            f"General Workers: {workers if workers is not None else 'Not specified'}",
            f"Supervisors: {supervisors if supervisors is not None else 'Not specified'}",
            f"Operational Hours: {hours_value} hours per day" if hours_value is not None else "Operational Hours: Not specified",
            f"Shifts: {shifts_value} per day" if shifts_value is not None else "Shifts: Not specified",
            f"Area Coverage: {area_value} m²" if area_value is not None else "Area Coverage: Not specified"
        ]
        
        for item in requirements_list:
            doc.add_paragraph(item, style='List Bullet')
    
    def _add_technical_proposal(self, doc, tender_data, pricing):
        req = tender_data.get('requirements', {})
        workforce = tender_data.get('workforce', {})
        
        doc.add_paragraph("Our proposed staffing allocation:", style='List Bullet')
        
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Light List Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Position'
        hdr_cells[1].text = 'Quantity'
        hdr_cells[2].text = 'Monthly Cost (R)'
        
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        labour_cost = pricing.get('labour_cost', 0) if isinstance(pricing, dict) else 0
        workers = workforce.get('total_workers')
        supervisors = workforce.get('supervisors')
        if workers is None:
            workers = req.get('workers')
        if supervisors is None:
            supervisors = req.get('supervisors')
        total_staff = safe_add(workers, supervisors)

        data = [
            ("General Workers", str(workers) if workers is not None else "Not specified", f"{labour_cost * 0.85:,.2f}"),
            ("Supervisors", str(supervisors) if supervisors is not None else "Not specified", f"{labour_cost * 0.15:,.2f}"),
            ("Total Labour", str(total_staff) if total_staff is not None else "Not specified", f"{labour_cost:,.2f}")
        ]
        
        for i, (pos, qty, cost) in enumerate(data, 1):
            row = table.rows[i]
            row.cells[0].text = pos
            row.cells[1].text = qty
            row.cells[2].text = cost
    
    def _add_pricing_schedule(self, doc, pricing):
        if not isinstance(pricing, dict):
            pricing = {}
        
        table = doc.add_table(rows=10, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr = table.rows[0].cells
        hdr[0].text = 'Item'
        hdr[1].text = 'Description'
        hdr[2].text = 'Amount (R)'
        
        for cell in hdr:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        items = [
            ('1', 'Labour Costs', pricing.get('labour_cost', 0)),
            ('2', 'Equipment', pricing.get('equipment_cost', 0)),
            ('3', 'Materials', pricing.get('materials_cost', 0)),
            ('4', 'Transport', pricing.get('transport_cost', 0)),
            ('', 'Subtotal', pricing.get('subtotal', 0)),
            ('', 'Overheads (15%)', pricing.get('overheads', 0)),
            ('', 'Profit (15%)', pricing.get('profit', 0)),
            ('', 'VAT (15%)', pricing.get('vat', 0)),
            ('', 'TOTAL MONTHLY', pricing.get('total_monthly', 0))
        ]
        
        for i, (num, desc, amt) in enumerate(items, 1):
            row = table.rows[i]
            row.cells[0].text = num
            row.cells[1].text = desc
            row.cells[2].text = f"{amt:,.2f}"
            
            if 'Subtotal' in desc or 'TOTAL' in desc:
                for cell in row.cells:
                    cell.paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        total_contract = pricing.get('total_contract_value', 0)
        run = para.add_run(f"Total Contract Value: R {total_contract:,.2f}")
        run.font.bold = True
        run.font.size = Pt(12)
        
        # Add pricing notes if available
        notes = pricing.get('notes')
        if notes:
            doc.add_paragraph()
            para = doc.add_paragraph()
            para.add_run("Pricing Notes:").font.bold = True
            doc.add_paragraph(notes)
    
    def _add_compliance_checklist(self, doc):
        table = doc.add_table(rows=6, cols=3)
        table.style = 'Table Grid'
        
        headers = ['Requirement', 'Status', 'Document Reference']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        
        requirements = [
            ("Tax Clearance Certificate", "☐ Attached", "SARS Pin"),
            ("B-BBEE Certificate", "☐ Attached", self.company_details.get('bbbee', '')),
            ("CIDB Registration", "☐ Attached", self.company_details.get('cidb', '')),
            ("Company Registration Docs", "☐ Attached", "CIPC Documents"),
            ("Workmans Compensation", "☐ Attached", "Letter of Good Standing")
        ]
        
        for i, (req, status, ref) in enumerate(requirements, 1):
            table.rows[i].cells[0].text = req
            table.rows[i].cells[1].text = status
            table.rows[i].cells[2].text = ref
    
    def _add_declaration(self, doc):
        content = """We hereby declare that:
1. The information furnished in this tender is true and correct.
2. We have not been involved in corrupt or fraudulent practices.
3. We accept the terms and conditions of this tender."""
        
        para = doc.add_paragraph(content.strip())
        para.paragraph_format.line_spacing = 1.15
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        table = doc.add_table(rows=2, cols=2)
        table.columns[0].width = Inches(3)
        table.columns[1].width = Inches(3)
        
        table.rows[0].cells[0].text = "_________________________"
        table.rows[0].cells[1].text = "_________________________"
        table.rows[1].cells[0].text = "Authorized Signature"
        table.rows[1].cells[1].text = "Date"
        
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    async def generate_excel(self, tender_data: Dict, pricing: Dict, output_path: str):
        loop = asyncio.get_event_loop()
        await loop.run_in_executor(None, self._create_formula_excel, tender_data, pricing, output_path)
    
    def _create_formula_excel(self, tender_data: Dict, pricing: Dict, output_path: str):
        wb = Workbook()
        
        # Ensure pricing is dict
        if not isinstance(pricing, dict):
            pricing = {}
        
        # Sheet 1: Labour
        ws1 = wb.active
        ws1.title = "Labour Costs"
        
        headers = ["Position", "Qty", "Hourly Rate", "Hours/Day", "Days/Month", "Monthly Rate", "Total Cost"]
        for col, header in enumerate(headers, 1):
            cell = ws1.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        
        req = tender_data.get('requirements', {})
        workforce = tender_data.get('workforce', {}) if isinstance(tender_data.get('workforce', {}), dict) else {}
        workers = workforce.get('total_workers')
        supervisors = workforce.get('supervisors')
        if workers is None:
            workers = req.get('workers')
        if supervisors is None:
            supervisors = req.get('supervisors')
        hours = req.get('hours_per_day')
        days = 22
        
        # Worker row
        ws1['A2'] = "General Workers"
        ws1['B2'] = workers
        ws1['C2'] = 25.00
        ws1['D2'] = hours
        ws1['E2'] = days
        ws1['F2'] = "=C2*D2*E2"
        ws1['G2'] = "=B2*F2"
        
        # Supervisor row
        ws1['A3'] = "Supervisors"
        ws1['B3'] = supervisors
        ws1['C3'] = 35.00
        ws1['D3'] = hours
        ws1['E3'] = days
        ws1['F3'] = "=C3*D3*E3"
        ws1['G3'] = "=B3*F3"
        
        ws1['A4'] = "Add: Deductions (35%)"
        ws1['G4'] = "=SUM(G2:G3)*0.35"
        
        ws1['A5'] = "TOTAL LABOUR"
        ws1['G5'] = "=SUM(G2:G3)+G4"
        ws1['G5'].font = Font(bold=True)
        
        # Sheet 2: Pricing
        ws2 = wb.create_sheet("Pricing Breakdown")
        
        ws2['A1'] = "Cost Component"
        ws2['B1'] = "Monthly Amount (R)"
        ws2['C1'] = "Formula"
        
        for cell in [ws2['A1'], ws2['B1'], ws2['C1']]:
            cell.font = Font(bold=True)
        
        row = 2
        components = [
            ("Labour Costs", "=Labour Costs!G5", "From labour sheet"),
            ("Equipment", pricing.get('equipment_cost', 3000), "Fixed"),
            ("Materials", pricing.get('materials_cost', 2000), "Fixed"),
            ("Transport", pricing.get('transport_cost', 2000), "Fixed")
        ]
        
        for name, value, formula in components:
            ws2.cell(row=row, column=1, value=name)
            ws2.cell(row=row, column=2, value=value)
            ws2.cell(row=row, column=3, value=formula)
            row += 1
        
        ws2[f'A{row}'] = "Subtotal"
        ws2[f'B{row}'] = f"=SUM(B2:B{row-1})"
        ws2[f'A{row}'].font = Font(bold=True)
        ws2[f'B{row}'].font = Font(bold=True)
        subtotal_row = row
        row += 1
        
        ws2[f'A{row}'] = "Overheads (15%)"
        ws2[f'B{row}'] = f"=B{subtotal_row}*0.15"
        row += 1
        
        ws2[f'A{row}'] = "Profit (15%)"
        ws2[f'B{row}'] = f"=B{subtotal_row}*0.15"
        row += 1
        
        ws2[f'A{row}'] = "VAT (15%)"
        ws2[f'B{row}'] = f"=(B{subtotal_row}+B{subtotal_row+1}+B{subtotal_row+2})*0.15"
        row += 1
        
        ws2[f'A{row}'] = "TOTAL MONTHLY"
        ws2[f'B{row}'] = f"=SUM(B{subtotal_row}:B{row-1})"
        ws2[f'A{row}'].font = Font(bold=True, size=12)
        ws2[f'B{row}'].font = Font(bold=True, size=12)
        
        # Auto-size
        for ws in [ws1, ws2]:
            for column in ws.columns:
                max_length = 0
                column_letter = get_column_letter(column[0].column)
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width
        
        wb.save(output_path)
