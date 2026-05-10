import re
import json
import logging
import pdfplumber
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
import asyncio
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
import io
import math
from datetime import datetime

logger = logging.getLogger(__name__)

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class ExtractionError(Exception):
    def __init__(self, payload: Dict[str, Any]):
        self.payload = payload
        super().__init__(str(payload))


class DocumentExtractor:
    """
    HONEST EXTRACTION - Multi-Sector Tender Intelligence
    Zero placeholders. Real estimates only when explicitly marked.
    """
    
    def __init__(self):
        self.extraction_confidence = {}
        self.estimated_values = {}
        # Internal estimator will be used for estimated contract value
        self.pdf_intel = None
    
    def _normalize_text(self, text: str) -> str:
        return ' '.join(text.strip().split())

    def _build_extracted_field(self, value: Optional[int], source: str, confidence: str, evidence: str) -> Dict[str, Any]:
        return {
            'value': value,
            'source': source,
            'confidence': confidence,
            'evidence': evidence.strip()[:260]
        }

    def _log_missing_fields(self, missing_fields: List[str], schedule: Dict[str, Any], sector: str, raw_text: str) -> None:
        snippet = raw_text[:1000].replace('\n', ' ').replace('\r', ' ')
        entry = {
            'timestamp': datetime.now().isoformat(),
            'missing_fields': missing_fields,
            'sector': sector,
            'detected_signals': schedule.get('detected_signals', []),
            'raw_text_snippet': snippet
        }
        logger.warning("Extraction missing fields: %s", entry)

    def _hours_context_valid(self, text: str) -> bool:
        return bool(re.search(r'\b(?:hour shift|hours per day|working hours|hrs/day|hrs per day|per shift|hours per shift|working hours per day)\b', text, re.IGNORECASE))

    def _reject_invalid_hours(self, hours: int, evidence: str) -> bool:
        if hours is None:
            return False
        if hours < 6:
            logger.debug("Rejected invalid hours (%s) evidence='%s'", hours, evidence)
            return True
        return False

    def _regex_schedule_extract(self, text: str, signals: List[str]) -> Dict[str, Optional[Dict[str, Any]]]:
        shifts = None
        hours = None

        patterns = [
            (r'(\d+)\s*[x×]\s*(\d{1,2})\s*hour\s*shifts?', 'explicit_x_hours_shifts'),
            (r'(\d+)\s*shifts?\s*(?:of\s*)?(\d{1,2})\s*hours?', 'explicit_shifts_of_hours'),
            (r'(\d+)\s*-?\s*hour\s*shift', 'hour_shift_phrase'),
            (r'(\d{1,2})\s*hours?\s*(?:per\s*day|per\s*shift|/day|/shift)\b', 'hours_per_shift'),
            (r'\bhours?\s*[:\-]\s*(\d{1,2})\b', 'hours_colon'),
            (r'(\d{1,2})\s*shifts?\b', 'shifts_count'),
            (r'shifts?\s*[:\-]?\s*(\d{1,2})', 'shifts_colon'),
        ]

        for pattern, label in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if not match:
                continue

            if label == 'explicit_x_hours_shifts' and len(match.groups()) >= 2:
                shifts = int(match.group(1))
                hours = int(match.group(2))
                evidence = match.group(0)
                if self._reject_invalid_hours(hours, evidence):
                    continue
                signals.append(f"{label}:{evidence}")
                return {
                    'shifts_per_day': self._build_extracted_field(shifts, 'explicit', 'high', evidence),
                    'hours_per_day': self._build_extracted_field(hours, 'explicit', 'high', evidence)
                }

            if label == 'explicit_shifts_of_hours' and len(match.groups()) >= 2:
                shifts = int(match.group(1))
                hours = int(match.group(2))
                evidence = match.group(0)
                if self._reject_invalid_hours(hours, evidence):
                    continue
                signals.append(f"{label}:{evidence}")
                return {
                    'shifts_per_day': self._build_extracted_field(shifts, 'explicit', 'high', evidence),
                    'hours_per_day': self._build_extracted_field(hours, 'explicit', 'high', evidence)
                }

            if label == 'hour_shift_phrase':
                hours = int(match.group(1))
                evidence = match.group(0)
                if self._reject_invalid_hours(hours, evidence):
                    continue
                signals.append(f"{label}:{evidence}")
                return {
                    'shifts_per_day': None,
                    'hours_per_day': self._build_extracted_field(hours, 'explicit', 'high', evidence)
                }

            if label == 'hours_per_shift' and hours is None:
                hours = int(match.group(1))
                evidence = match.group(0)
                if self._reject_invalid_hours(hours, evidence):
                    continue
                signals.append(f"{label}:{evidence}")
                return {
                    'shifts_per_day': None,
                    'hours_per_day': self._build_extracted_field(hours, 'explicit', 'high', evidence)
                }

            if label in {'shifts_count', 'shifts_colon'} and shifts is None:
                shifts = int(match.group(1))
                signals.append(f"{label}:{match.group(0)}")
                return {
                    'shifts_per_day': self._build_extracted_field(shifts, 'explicit', 'high', match.group(0)),
                    'hours_per_day': None
                }

            if label == 'hours_colon' and hours is None:
                hours = int(match.group(1))
                evidence = match.group(0)
                if self._reject_invalid_hours(hours, evidence):
                    continue
                signals.append(f"{label}:{evidence}")
                return {
                    'shifts_per_day': None,
                    'hours_per_day': self._build_extracted_field(hours, 'explicit', 'high', evidence)
                }

        return {'shifts_per_day': None, 'hours_per_day': None}

    def _semantic_schedule_infer(self, text: str, signals: List[str]) -> Dict[str, Optional[Dict[str, Any]]]:
        lowered = text.lower()

        if '24-hour operation' in lowered or '24 hour operation' in lowered or '24-hour operations' in lowered:
            evidence = self._find_evidence(text, r'24[- ]hour operation')
            signals.append(f"semantic:24_hour_operation:{evidence}")
            return {
                'shifts_per_day': self._build_extracted_field(3, 'inferred', 'high', evidence),
                'hours_per_day': self._build_extracted_field(8, 'inferred', 'high', evidence)
            }

        if 'day and night shift' in lowered or 'day-night shift' in lowered or 'day and night shifts' in lowered:
            evidence = self._find_evidence(text, r'(?:day and night shift|day-night shift|day and night shifts)')
            signals.append(f"semantic:day_and_night:{evidence}")
            return {
                'shifts_per_day': self._build_extracted_field(2, 'inferred', 'medium', evidence),
                'hours_per_day': self._build_extracted_field(12, 'inferred', 'medium', evidence)
            }

        time_range = re.search(r'(\d{1,2}:\d{2})\s*[–-]\s*(\d{1,2}:\d{2})', text)
        if time_range:
            start = time_range.group(1)
            end = time_range.group(2)
            evidence = time_range.group(0)
            duration = self._infer_duration_from_time_range(start, end)
            if duration is not None:
                signals.append(f"semantic:time_range:{evidence}")
                return {
                    'shifts_per_day': self._build_extracted_field(1, 'inferred', 'high', evidence),
                    'hours_per_day': self._build_extracted_field(duration, 'inferred', 'high', evidence)
                }

        return {'shifts_per_day': None, 'hours_per_day': None}

    def _contextual_schedule_infer(self, text: str, signals: List[str]) -> Dict[str, Optional[Dict[str, Any]]]:
        lowered = text.lower()

        total_match = re.search(r'(\d+)\s*(?:total staff|total workers|total employees|workers in total)', text, re.IGNORECASE)
        per_shift_match = re.search(r'(\d+)\s*(?:per shift|/shift)', text, re.IGNORECASE)
        if total_match and per_shift_match:
            total = int(total_match.group(1))
            per_shift = int(per_shift_match.group(1))
            if per_shift > 0 and total % per_shift == 0:
                shifts = total // per_shift
                evidence = f"{total_match.group(0)} / {per_shift_match.group(0)}"
                signals.append(f"contextual:total_per_shift:{evidence}")
                return {
                    'shifts_per_day': self._build_extracted_field(shifts, 'inferred', 'medium', evidence),
                    'hours_per_day': None
                }

        if 'per shift' in lowered and 'hour' in lowered:
            evidence = self._find_evidence(text, r'\d+\s*hours?\s*per\s*shift|per shift')
            signals.append(f"contextual:per_shift_hours:{evidence}")
            hour_match = re.search(r'(\d+)\s*hours?\s*per\s*shift', text, re.IGNORECASE)
            if hour_match:
                hours = int(hour_match.group(1))
                return {'shifts_per_day': None, 'hours_per_day': self._build_extracted_field(hours, 'inferred', 'medium', hour_match.group(0))}

        return {'shifts_per_day': None, 'hours_per_day': None}

    def _infer_duration_from_time_range(self, start: str, end: str) -> Optional[int]:
        try:
            start_h, start_m = map(int, start.split(':'))
            end_h, end_m = map(int, end.split(':'))
            start_minutes = start_h * 60 + start_m
            end_minutes = end_h * 60 + end_m
            if end_minutes <= start_minutes:
                end_minutes += 24 * 60
            duration = end_minutes - start_minutes
            if duration % 60 == 0:
                return duration // 60
        except ValueError:
            return None
        return None

    def _find_evidence(self, text: str, pattern: str) -> str:
        match = re.search(pattern, text, re.IGNORECASE)
        return match.group(0) if match else pattern

    def _detect_24h_coverage(self, text: str) -> Optional[str]:
        lowered = text.lower()
        patterns = [
            r'24\s*-?\s*hour\s*shift',
            r'24\s*-?\s*hour\s*service',
            r'24\s*-?\s*hour\s*coverage',
            r'round\s*the\s*clock',
            r'continuous\s*operation'
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(0)
        return None

    def _normalize_24h_schedule(self, schedule: Dict[str, Any], text: str, signals: List[str]) -> Dict[str, Any]:
        hours_field = schedule.get('hours_per_day')
        shifts_field = schedule.get('shifts_per_day')
        if hours_field and hours_field.get('value') == 24:
            coverage_evidence = self._detect_24h_coverage(text)
            if coverage_evidence and not shifts_field:
                logger.debug("Detected 24-hour coverage pattern")
                shift_evidence = "24 Hour Shift → interpreted as 24h coverage"
                schedule['shifts_per_day'] = self._build_extracted_field(3, 'inferred', 'high', shift_evidence)
                schedule['hours_per_day'] = self._build_extracted_field(8, 'inferred', 'high', "24 Hour Shift → standard 3x8h shift model")
                signals.append(f"semantic:24h_coverage:{coverage_evidence}")
                logger.debug("Converted to 3 shifts x 8 hours")
                return schedule
            logger.debug("Invalid hours_per_day == 24 without valid 24h coverage pattern")
            schedule['hours_per_day'] = None
            return schedule

        if hours_field and hours_field.get('value', 0) > 12:
            logger.debug("Invalid hours_per_day > 12 detected")
            schedule['hours_per_day'] = None
        return schedule

    def _apply_sector_schedule_rules(self, schedule: Dict[str, Any], text: str, sector: str, signals: List[str]) -> Dict[str, Any]:
        lowered = text.lower()

        if sector == 'security' and ('24 hour' in lowered or '24-hour' in lowered or 'round the clock' in lowered):
            evidence = self._find_evidence(text, r'24[- ]hour|round the clock')
            signals.append(f"sector:security_24hour:{evidence}")
            if not schedule['shifts_per_day']:
                schedule['shifts_per_day'] = self._build_extracted_field(3, 'inferred', 'high', evidence)
            if not schedule['hours_per_day']:
                schedule['hours_per_day'] = self._build_extracted_field(8, 'inferred', 'high', evidence)

        if sector == 'cleaning' and not schedule['shifts_per_day']:
            if 'single shift' in lowered or 'one shift' in lowered:
                evidence = self._find_evidence(text, r'(single shift|one shift)')
                signals.append(f"sector:cleaning_single_shift:{evidence}")
                schedule['shifts_per_day'] = self._build_extracted_field(1, 'inferred', 'medium', evidence)
            elif 'two shift' in lowered or 'two shifts' in lowered:
                evidence = self._find_evidence(text, r'(two shift|two shifts)')
                signals.append(f"sector:cleaning_two_shift:{evidence}")
                schedule['shifts_per_day'] = self._build_extracted_field(2, 'inferred', 'medium', evidence)

        if sector == 'construction' and not schedule['hours_per_day']:
            match = re.search(r'(?:8|10)\s*-?\s*hour\s*day', text, re.IGNORECASE)
            if match:
                hours = int(re.search(r'(8|10)', match.group(0)).group(1))
                evidence = match.group(0)
                signals.append(f"sector:construction_hours:{evidence}")
                schedule['hours_per_day'] = self._build_extracted_field(hours, 'inferred', 'high', evidence)

        return schedule

    def extract_work_schedule(self, text: str, sector: Optional[str] = None) -> Dict[str, Any]:
        signals: List[str] = []

        explicit = self._regex_schedule_extract(text, signals)
        schedule = {
            'shifts_per_day': explicit['shifts_per_day'],
            'hours_per_day': explicit['hours_per_day'],
            'detected_signals': signals
        }

        logger.debug("Candidate explicit schedule shifts=%s hours=%s signals=%s", schedule['shifts_per_day'], schedule['hours_per_day'], signals)
        if schedule['shifts_per_day'] and schedule['hours_per_day']:
            if sector:
                schedule = self._apply_sector_schedule_rules(schedule, text, sector, signals)
            return schedule

        semantic = self._semantic_schedule_infer(text, signals)
        if semantic['shifts_per_day'] and not schedule['shifts_per_day']:
            schedule['shifts_per_day'] = semantic['shifts_per_day']
        if semantic['hours_per_day'] and not schedule['hours_per_day']:
            schedule['hours_per_day'] = semantic['hours_per_day']

        if schedule['shifts_per_day'] and schedule['hours_per_day']:
            if sector:
                schedule = self._apply_sector_schedule_rules(schedule, text, sector, signals)
            return schedule

        contextual = self._contextual_schedule_infer(text, signals)
        if contextual['shifts_per_day'] and not schedule['shifts_per_day']:
            schedule['shifts_per_day'] = contextual['shifts_per_day']
        if contextual['hours_per_day'] and not schedule['hours_per_day']:
            schedule['hours_per_day'] = contextual['hours_per_day']

        if sector:
            schedule = self._apply_sector_schedule_rules(schedule, text, sector, signals)

        schedule = self._normalize_24h_schedule(schedule, text, signals)

        logger.debug("Final schedule after inference shifts=%s hours=%s signals=%s", schedule['shifts_per_day'], schedule['hours_per_day'], signals)
        return schedule

    def _estimate_value_from_text(self, text: str) -> Dict[str, Any]:
        """Attempt to extract an estimated contract value from text.

        Returns a dict with keys: estimated_value (float|None), evidence, source, confidence
        """
        if not text:
            return {'estimated_value': None, 'evidence': None, 'source': None, 'confidence': 'low'}

        # Common currency pattern e.g. R 1,234,567.89
        m = re.search(r'R\s?([\d,]+(?:\.\d+)?)', text, re.IGNORECASE)
        if m:
            num = m.group(1).replace(',', '')
            try:
                val = float(num)
                return {'estimated_value': round(val, 2), 'evidence': m.group(0), 'source': 'text_regex', 'confidence': 'medium'}
            except Exception:
                pass

        # e.g. 1.2 million
        m = re.search(r'([\d,]+(?:\.\d+)?)\s*(million|m)\b', text, re.IGNORECASE)
        if m:
            num = m.group(1).replace(',', '')
            try:
                val = float(num) * 1_000_000
                return {'estimated_value': round(val, 2), 'evidence': m.group(0), 'source': 'text_regex', 'confidence': 'low'}
            except Exception:
                pass

        # fallback: no reliable value found
        return {'estimated_value': None, 'evidence': None, 'source': None, 'confidence': 'low'}

    async def extract(self, file_path: str) -> Dict:
        """Extract with sector-specific intelligence"""
        logger.debug("[EXTRACT] Started")
        path = Path(file_path)
        ext = path.suffix.lower()
        logger.debug("File path=%s ext=%s", file_path, ext)
        
        # Extract raw text
        logger.debug("Running raw text extraction")
        if ext == '.pdf':
            text = await self._extract_pdf(path)
        elif ext in ['.docx', '.doc']:
            text = await self._extract_docx(path)
        else:
            text = await self._extract_text(path)
        logger.debug("Raw text extracted")
        
        # KEEP FULL TEXT for comprehensive search
        full_text = self._normalize_text(text)
        logger.debug("Normalized text")
        text_upper = full_text.upper()
        
        # DETECT SECTOR
        sector_data = self._detect_sector_and_type(text_upper)
        sector = sector_data['sector']
        sub_type = sector_data['sub_type']
        logger.debug("Detected sector=%s sub_type=%s", sector, sub_type)
        
        # Extract duration with context
        duration_data = self._extract_duration_intelligent(text_upper, sector)
        logger.debug("Duration data=%s", duration_data)
        
        # Extract or estimate workers using FULL TEXT
        worker_data = self._estimate_workers_by_sector(text_upper, full_text, sector, duration_data.get('months') if isinstance(duration_data, dict) else None)
        logger.debug("Worker data=%s", worker_data)
        
        # Extract or estimate area/scope
        scope_data = self._estimate_scope_by_sector(text_upper, full_text, sector)
        logger.debug("Scope data=%s", scope_data)
        
        # Extract location
        location = self._extract_location(text_upper)
        logger.debug("Location=%s", location)

        # Schedule extraction pipeline: regex -> semantic -> contextual -> sector rules
        schedule = self.extract_work_schedule(full_text, sector)
        logger.debug("Schedule data=%s", schedule)
        shifts_field = schedule['shifts_per_day']
        hours_field = schedule['hours_per_day']
        shifts = shifts_field['value'] if shifts_field else None
        hours = hours_field['value'] if hours_field else None
        
        # Estimate contract value from text (internal)
        intel_data = self._estimate_value_from_text(full_text)
        estimated_value = intel_data.get('estimated_value')
        
        # Build result
        missing_fields = []
        total_workers = worker_data.get('total')
        if total_workers is None:
            missing_fields.append('workers')
        supervisors = worker_data.get('supervisors')
        duration_months = duration_data.get('months')
        area_sqm = scope_data.get('area')
        working_days = self._extract_working_days(text_upper)
        result = {
            'reference': self._extract_reference(text_upper),
            'closing_date': self._extract_closing_date(full_text),
            'sector': sector,
            'sub_type': sub_type,
            'location': location,
            'estimated_value': estimated_value,

            'duration': duration_data,

            'scope': {
                'area_sqm': area_sqm,
                'area_note': scope_data.get('note'),
                'description': self._extract_scope_description(full_text, sector),
                'is_emergency': 'EMERGENCY' in text_upper or 'URGENT' in text_upper,
                'is_breakdown': 'BREAKDOWN' in text_upper or 'REPAIR' in text_upper
            },

            'shifts_per_day': shifts,
            'hours_per_day': hours,
            'schedule': schedule,

            'workforce': {
                'total_workers': total_workers,
                'skilled_workers': worker_data.get('skilled'),
                'unskilled_workers': worker_data.get('unskilled'),
                'supervisors': supervisors,
                'estimation_method': worker_data.get('method'),
                'confidence': worker_data.get('confidence')
            },

            'workers': {
                'total': total_workers,
                'supervisors': supervisors
            },
            'missing_fields': missing_fields,

            'requirements': {
                'shifts_per_day': shifts,
                'hours_per_day': hours,
                'schedule': schedule,
                'working_days_week': working_days,
                'equipment_required': self._extract_equipment_by_sector(text_upper, sector),
                'materials_required': self._extract_materials_by_sector(text_upper, sector),
                'certifications_required': self._extract_certifications(text_upper, sector),
                # Backward-compat fields for document_generator, debate_council, index.html
                'workers': total_workers,
                'supervisors': supervisors,
                'area_sqm': area_sqm,
                'shifts': shifts,
                'total_staff': None if total_workers is None or supervisors is None else total_workers + supervisors
            },

            'client': {
                'type': self._extract_client_type(text_upper),
                'name': self._extract_client_name(full_text)
            },

            '_extraction_notes': {
                'sector_detected': sector,
                'worker_estimated': worker_data.get('estimated', False),
                'duration_estimated': duration_data.get('estimated', False),
                'area_estimated': scope_data.get('estimated', False),
                'raw_confidence': self._calculate_overall_confidence(worker_data, duration_data, scope_data)
            },

            'raw_text': full_text[:5000],

            # === BACKWARD-COMPAT TOP-LEVEL FIELDS ===
            # These duplicate nested data so downstream code using the old flat format still works
            'tender_type': sector,
            'duration_months': duration_months
        }

        missing = [f for f in ('shifts_per_day', 'hours_per_day') if result.get(f) is None]

        if missing:
            payload = {
                'status': 'incomplete',
                'missing_fields': missing,
                'extracted_data': {
                    'shifts_per_day': shifts_field,
                    'hours_per_day': hours_field
                },
                'detected_signals': schedule.get('detected_signals', []) if isinstance(schedule, dict) else []
            }
            # Build suggested inputs derived from signals (no guessing)
            suggested = {}
            for s in payload.get('detected_signals', []):
                s_low = s.lower()
                if '24_hour' in s_low or '24h' in s_low:
                    suggested.setdefault('shifts_per_day', []).append(3)
                    suggested.setdefault('hours_per_day', []).append(8)
                if 'day_and_night' in s_low or 'day-night' in s_low:
                    suggested.setdefault('shifts_per_day', []).append(2)
                    suggested.setdefault('hours_per_day', []).append(12)
            if suggested:
                payload['suggested_inputs'] = suggested

            # Log structured issue and raise
            self._log_missing_fields(missing, schedule, sector, full_text)
            raise ExtractionError(payload)

        return result
    
    def _detect_sector_and_type(self, text: str) -> Dict:
        """Detect sector with weighted scoring - FIXED for cleaning detection"""
        
        # CRITICAL FIX #1: Check for exact phrase "CLEANING SERVICES" first (high priority)
        if 'CLEANING SERVICES' in text:
            return {'sector': 'cleaning', 'sub_type': 'commercial', 'confidence': 15}
        
        # Also check common cleaning patterns
        cleaning_phrases = ['OFFICE CLEANING', 'CLEANING CONTRACT', 'JANITORIAL SERVICES', 
                           'HYGIENE SERVICES', 'CLEANING AND', 'CLEANING OF']
        for phrase in cleaning_phrases:
            if phrase in text:
                return {'sector': 'cleaning', 'sub_type': 'standard', 'confidence': 12}
        
        # CONSTRUCTION
        construction_terms = {
            'construction': 5, 'building': 5, 'civil': 5, 'structural': 4,
            'concrete': 4, 'brickwork': 4, 'plastering': 4, 'roofing': 4,
            'excavation': 4, 'foundation': 4, 'renovation': 3, 'refurbishment': 3,
            'plumbing': 3, 'carpentry': 3, 'painting': 3, 'paving': 3,
            'asphalt': 4, 'tar': 3, 'road': 3, 'bridge': 4,
            'demolition': 3, 'earthworks': 4, 'steel': 3, 'formwork': 4,
            'scaffolding': 3, 'drainage': 3, 'sewer': 3, 'waterproofing': 3
        }
        scores = {'construction': sum(weight for term, weight in construction_terms.items() if term in text)}
        
        # ELECTRICAL
        electrical_terms = {
            'electrical': 5, 'electrician': 5, 'wiring': 4, 'cable': 3,
            'distribution board': 5, 'substation': 4, 'high voltage': 5,
            'solar': 4, 'generator': 3, 'reticulation': 3, 'transformer': 4
        }
        scores['electrical'] = sum(weight for term, weight in electrical_terms.items() if term in text)
        
        # CLEANING (individual terms - lower weight than phrases above)
        cleaning_terms = {
            'cleaning': 3, 'cleaner': 3, 'hygiene': 4, 'sanitiz': 4,
            'janitorial': 4, 'housekeeping': 3, 'deep clean': 4
        }
        scores['cleaning'] = sum(weight for term, weight in cleaning_terms.items() if term in text)
        
        # SECURITY
        security_terms = {
            'security': 5, 'guarding': 5, 'cctv': 4, 'surveillance': 4,
            'access control': 5, 'alarm': 3, 'patrol': 4, 'psira': 5
        }
        scores['security'] = sum(weight for term, weight in security_terms.items() if term in text)
        
        # GARDENING
        garden_terms = {
            'garden': 5, 'landscaping': 5, 'grass': 3, 'lawn': 3,
            'tree': 3, 'felling': 4, 'pruning': 4, 'irrigation': 4,
            'grounds': 4, 'horticulture': 5
        }
        scores['gardening'] = sum(weight for term, weight in garden_terms.items() if term in text)
        
        # IT
        it_terms = {
            'software': 5, 'programming': 5, 'development': 4, 'it services': 5,
            'network': 4, 'server': 4, 'cloud': 4, 'cyber': 4,
            'database': 4, 'application': 3, 'helpdesk': 4,
            'coding': 5, 'java': 4, 'python': 4, 'sql': 4
        }
        scores['it_services'] = sum(weight for term, weight in it_terms.items() if term in text)
        
        # MAINTENANCE
        maintenance_terms = {
            'maintenance': 4, 'repair': 3, 'service': 2, 'breakdown': 4
        }
        scores['maintenance'] = sum(weight for term, weight in maintenance_terms.items() if term in text)
        
        # SUPPLY
        supply_terms = {
            'supply': 5, 'deliver': 4, 'procure': 4, 'materials': 3,
            'equipment': 3, 'goods': 3, 'tools': 3
        }
        scores['supply'] = sum(weight for term, weight in supply_terms.items() if term in text)
        
        # Determine winner
        best_sector = max(scores, key=scores.get)
        
        # Handle low confidence
        if scores[best_sector] < 5:
            best_sector = 'general'
            confidence = 0
        else:
            confidence = scores[best_sector]
        
        # Detect sub-type
        sub_type = 'standard'
        if 'EMERGENCY' in text or 'URGENT' in text:
            sub_type = 'emergency'
        elif 'RESIDENTIAL' in text or 'HOUSING' in text:
            sub_type = 'residential'
        elif 'INDUSTRIAL' in text or 'FACTORY' in text:
            sub_type = 'industrial'
        elif 'COMMERCIAL' in text or 'OFFICE' in text:
            sub_type = 'commercial'
        elif 'REPAIR' in text or 'BREAKDOWN' in text:
            sub_type = 'repair'
        
        return {'sector': best_sector, 'sub_type': sub_type, 'confidence': confidence}
    
    def _estimate_workers_by_sector(self, text: str, raw_text: str, sector: str, duration_months: int) -> Dict:
        """HONEST ESTIMATION: Search full document for workers"""
        
        # Search PATTERNS in priority order
        patterns = [
            r'(\d{1,3})\s*CLEANERS?',
            r'(\d{1,3})\s*WORKERS?',
            r'(\d{1,3})\s*STAFF',
            r'(\d{1,3})\s*PERSONNEL',
            r'(\d{1,3})\s*EMPLOYEES?',
            r'(\d{1,3})\s*OPERATIVES?',
            r'(\d{1,3})\s*LABOURERS?',
            r'(\d{1,3})\s*PEOPLE',
            r'TEAM\s*OF\s*(\d{1,3})',
            r'NO[.:]?\s*OF\s*CLEANERS?\s*:?\s*(\d{1,3})',
            r'NO[.:]?\s*OF\s*WORKERS?\s*:?\s*(\d{1,3})',
            r'QUANTITY[:\s]+(\d{1,3})\s*(?:CLEANERS?|WORKERS?)',
            r'(\d{1,3})\s*DAY\s*SHIFT',
            r'(\d{1,3})\s*NIGHT\s*SHIFT'
        ]
        
        # Search in FULL raw_text (not just first 1000 chars)
        for pattern in patterns:
            matches = re.findall(pattern, raw_text.upper())
            if matches:
                try:
                    # Handle both direct match and group match
                    match = matches[0]
                    workers = int(match) if isinstance(match, str) else int(match[0])
                    if 1 <= workers <= 500:
                        return {
                            'total': workers,
                            'method': 'Extracted from document',
                            'estimated': False,
                            'confidence': 'High',
                            'note': f'Found: {workers} workers'
                        }
                except:
                    continue
        
        return {'estimated': True, 'note': 'Unable to estimate workforce from document'}
    
    def _extract_duration_intelligent(self, text: str, sector: str) -> Dict:
        """FIXED: Look for contract period, not RFQ validity"""
        
        # PRIORITY 1: Look for "period of twelve (12) months" or similar
        contract_patterns = [
            r'period\s+of\s+(\w+)\s+\((\d{1,2})\)\s*months?',  # "twelve (12) months"
            r'period\s+of\s+(\d{1,2})\s*months?',  # "12 months"
            r'agreement.*?for\s+(\d{1,2})\s*months',  # "agreement for 12 months"
            r'duration.*?(\d{1,2})\s*months',  # "duration: 12 months"
            r'contract.*?(\d{1,2})\s*months',  # "contract period 12 months"
            r'for\s+a\s+period\s+of\s+(\d{1,2})\s*months',  # "for a period of 12 months"
            r'twelve\s*\(12\)\s*months',  # "twelve (12) months"
            r'\(12\)\s*months',  # "(12) months"
        ]
        
        for pattern in contract_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                try:
                    months = int(match.group(2)) if len(match.groups()) > 1 and match.group(2) else int(match.group(1))
                except:
                    continue
                
                return {
                    'value': months,
                    'unit': 'months',
                    'months': months,
                    'display': f'{months} months',
                    'estimated': False,
                    'note': 'Extracted from contract period'
                }
        
        # Check for written "TWELVE"
        if 'TWELVE' in text and 'MONTH' in text:
            return {
                'value': 12,
                'unit': 'months',
                'months': 12,
                'display': '12 months',
                'estimated': False,
                'note': 'Extracted: "twelve months"'
            }
        
        # Years
        year_match = re.search(r'(\d)\s*YEAR', text)
        if year_match:
            years = int(year_match.group(1))
            return {
                'value': years,
                'unit': 'years',
                'months': years * 12,
                'display': f'{years} years',
                'estimated': False
            }
        
        # Emergency days
        if 'EMERGENCY' in text:
            day_match = re.search(r'(\d{1,3})\s*DAYS?', text)
            if day_match:
                days = int(day_match.group(1))
                return {
                    'value': days,
                    'unit': 'days',
                    'months': days / 30,
                    'display': f'{days} days',
                    'estimated': False,
                    'is_emergency': True
                }
        
        return {'estimated': True, 'note': 'Unable to determine duration from document'}
    
    def _estimate_scope_by_sector(self, text: str, raw_text: str, sector: str) -> Dict:
        """Estimate scope"""
        area_match = re.search(r'(\d{1,6}(?:,\d{3})?)\s*(?:M2|M²|SQM|SQUARE)', text)
        if area_match:
            return {
                'area': int(area_match.group(1).replace(',', '')),
                'estimated': False,
                'note': 'Extracted from document'
            }
        
        return {'estimated': True, 'note': 'Unable to determine scope from document'}
    
    def _extract_equipment_by_sector(self, text: str, sector: str) -> List[str]:
        text_lower = text.lower()
        equipment = []
        
        if sector == 'cleaning':
            items = ['scrubber', 'vacuum', 'polisher', 'pressure washer', 'carpet cleaner']
            for item in items:
                if item in text_lower:
                    equipment.append(item)
        
        elif sector == 'construction':
            items = ['excavator', 'tipper', 'mixer', 'scaffolding', 'compactor', 'crane']
            for item in items:
                if item in text_lower:
                    equipment.append(item)
        
        return equipment if equipment else []
    
    def _extract_materials_by_sector(self, text: str, sector: str) -> List[str]:
        text_lower = text.lower()
        materials = []
        
        if sector == 'cleaning':
            items = ['detergent', 'disinfectant', 'cloths', 'paper towels']
        elif sector == 'construction':
            items = ['cement', 'bricks', 'steel', 'sand']
        elif sector == 'electrical':
            items = ['cable', 'distribution board', 'breakers']
        else:
            return []
        
        for item in items:
            if item in text_lower:
                materials.append(item)
        
        return materials
    
    def _extract_certifications(self, text: str, sector: str) -> List[str]:
        certs = []
        if sector == 'electrical':
            certs.append('Master Electrician Certificate')
        elif sector == 'security':
            certs.append('PSIRA Registration')
        elif sector == 'cleaning':
            certs.append('Cleaning Industry Registration')
        elif sector == 'construction':
            certs.append('CIDB Registration')
        return certs
    
    def _extract_working_days(self, text: str) -> Optional[int]:
        if '7 DAYS' in text:
            return 7
        if '6 DAYS' in text:
            return 6
        if '5 DAYS' in text:
            return 5
        return None
    
    def _extract_location(self, text: str) -> Optional[str]:
        cities = ['johannesburg', 'pretoria', 'cape town', 'durban', 'port elizabeth', 
                 'east london', 'bloemfontein', 'nelspruit', 'polokwane', 'george', 
                 'plettenberg bay', 'east london']
        
        text_lower = text.lower()
        for city in cities:
            if city in text_lower:
                return city.title()
        return None
    
    def _extract_reference(self, text: str) -> Optional[str]:
        patterns = [
            r'RFQ\s*[No.]*\s*(\d{6,12})',
            r'PR\s*(\d{6,12})',
            r'TENDER\s*NO[.:]\s*(\w{5,20})'
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(0).replace(' ', '').replace(':', '')[:30]
        return None
    
    def _extract_closing_date(self, text: str) -> Optional[str]:
        match = re.search(r'(\d{1,2}[\/\-.]\d{1,2}[\/\-.]\d{2,4})', text)
        if match:
            return match.group(1)
        return None
    
    def _extract_scope_description(self, text: str, sector: str) -> Optional[str]:
        lines = text.split('\n')
        for line in lines:
            clean = line.strip()
            if len(clean) > 50 and any(word in clean.lower() for word in ['scope', 'description', 'background']):
                return clean[:200]
        return None
    
    def _extract_client_type(self, text: str) -> Optional[str]:
        if 'MUNICIPALITY' in text:
            return 'Municipality'
        elif 'GOVERNMENT' in text:
            return 'Government'
        elif 'SCHOOL' in text:
            return 'Educational'
        return None
    
    def _extract_client_name(self, text: str) -> Optional[str]:
        match = re.search(r'(?:FROM|BY)\s*([A-Z][A-Za-z\s]{10,50})', text)
        if match:
            return match.group(1).strip()[:50]
        return None
    
    def _calculate_overall_confidence(self, worker_data, duration_data, scope_data) -> str:
        scores = []
        if not worker_data.get('estimated'):
            scores.append(1.0)
        else:
            scores.append(0.5 if worker_data.get('confidence') == 'Medium' else 0.3)
        if not duration_data.get('estimated'):
            scores.append(1.0)
        else:
            scores.append(0.7)
        if not scope_data.get('estimated'):
            scores.append(1.0)
        else:
            scores.append(0.6)
        
        avg = sum(scores) / len(scores)
        if avg >= 0.8:
            return "High"
        elif avg >= 0.5:
            return "Medium"
        else:
            return "Low - Verify"
    
    async def _extract_pdf(self, pdf_path: Path) -> str:
        text = ""
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.exception("pdfplumber error: %s", e)
        if len(text.strip()) < 500:
            text = await self._ocr_pdf(pdf_path)
        return text
    
    async def _ocr_pdf(self, pdf_path: Path) -> str:
        text = ""
        try:
            images = convert_from_path(str(pdf_path), dpi=150)
            for i, image in enumerate(images[:20]):  # INCREASED to 20 pages
                page_text = pytesseract.image_to_string(image)
                text += f"\nPAGE {i+1}:\n{page_text}\n"
        except Exception as e:
            logger.exception("OCR error: %s", e)
        return text
    
    async def _extract_docx(self, docx_path: Path) -> str:
        if not DOCX_AVAILABLE:
            try:
                import zipfile
                with zipfile.ZipFile(docx_path) as zf:
                    xml_content = zf.read('word/document.xml')
                    text = re.sub(r'<[^>]+>', '', xml_content.decode('utf-8', errors='ignore'))
                    return text
            except:
                return ""
        try:
            doc = Document(str(docx_path))
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text:
                        full_text.append(row_text)
            return '\n'.join(full_text)
        except:
            return ""
    
    async def _extract_text(self, txt_path: Path) -> str:
        try:
            with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except:
            return ""

# Backward compatibility
class PDFProcessor(DocumentExtractor):
    pass
